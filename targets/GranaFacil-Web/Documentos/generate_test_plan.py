#!/usr/bin/env python3
"""
Gerador de Plano de Teste (Test Plan) - Banco Digital Grana Fácil
Template profissional para documentação de pentest
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, timedelta

def add_cover_page(doc):
    """Adiciona capa do plano de teste"""
    # Logo/Título centralizado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")
    
    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLANO DE TESTE DE INTRUSÃO\n")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(Test Plan)")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(118, 75, 162)
    
    # Empresa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nBanco Digital Grana Fácil")
    run.font.size = Pt(20)
    run.font.bold = True
    
    # Tipo de teste
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nPentest Web - Black Box")
    run.font.size = Pt(16)
    
    # Classificação
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nCONFIDENCIAL - INTERNAL USE ONLY")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    # Espaçamento
    doc.add_paragraph("\n\n\n")
    
    # Informações do teste
    hoje = date.today()
    fim = hoje + timedelta(days=14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Período de Teste: {hoje.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}\n")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versão: 1.0")
    run.font.size = Pt(12)
    
    # Quebra de página
    doc.add_page_break()

def add_document_info(doc):
    """Informações do documento"""
    doc.add_heading('INFORMAÇÕES DO DOCUMENTO', 1)
    
    hoje = date.today()
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Título', 'Plano de Teste de Intrusão - Banco Digital Grana Fácil'),
        ('Tipo de Teste', 'Pentest Web - Black Box'),
        ('Versão', '1.0'),
        ('Data de Criação', hoje.strftime('%d/%m/%Y')),
        ('Preparado por', 'AulasHack Security Team'),
        ('Aprovado por', 'João Silva - CISO'),
        ('Classificação', 'CONFIDENCIAL'),
        ('Status', 'Aprovado')
    ]
    
    for i, (label, value) in enumerate(info_data):
        row_cells = table.rows[i].cells
        row_cells[0].text = label
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = value
    
    doc.add_page_break()

def add_toc(doc):
    """Sumário"""
    doc.add_heading('SUMÁRIO', 1)
    
    sections = [
        '1. INFORMAÇÕES DO PROJETO',
        '2. CONTATOS',
        '   2.1 Equipe de Pentest',
        '   2.2 Contatos do Cliente',
        '3. ESCOPO DO TESTE',
        '   3.1 Objetivos',
        '   3.2 Alvos no Escopo',
        '   3.3 Fora do Escopo',
        '4. TIPO E METODOLOGIA',
        '   4.1 Tipo de Teste: Black Box',
        '   4.2 Metodologia',
        '5. CRONOGRAMA',
        '6. LIMITAÇÕES E RESTRIÇÕES',
        '   6.1 Restrições de Horário',
        '   6.2 Atividades Proibidas',
        '   6.3 Limitações Técnicas',
        '7. REGRAS DE ENGAJAMENTO',
        '   7.1 Comunicação',
        '   7.2 Procedimentos de Emergência',
        '   7.3 Confidencialidade',
        '8. ENTREGÁVEIS',
        '9. CRITÉRIOS DE ACEITAÇÃO',
        '10. APROVAÇÕES'
    ]
    
    for section in sections:
        p = doc.add_paragraph(section)
        if not section.startswith('   '):
            p.style = 'List Number'
        else:
            p.style = 'List Bullet'
    
    doc.add_page_break()

def add_section1(doc):
    """Seção 1: Informações do Projeto"""
    doc.add_heading('1. INFORMAÇÕES DO PROJETO', 1)
    
    doc.add_heading('1.1 Nome do Projeto', 2)
    doc.add_paragraph(
        "Teste de Intrusão da Aplicação Web do Banco Digital Grana Fácil"
    )
    
    doc.add_heading('1.2 Cliente', 2)
    doc.add_paragraph(
        "Banco Digital Grana Fácil S.A.\n"
        "CNPJ: 12.345.678/0001-90\n"
        "Endereço: Av. Paulista, 1000 - São Paulo/SP\n"
        "CEP: 01310-100"
    )
    
    doc.add_heading('1.3 Empresa Executora', 2)
    doc.add_paragraph(
        "AulasHack Security Consulting\n"
        "CNPJ: 98.765.432/0001-10\n"
        "Endereço: Rua da Segurança, 500 - São Paulo/SP\n"
        "CEP: 04567-890\n"
        "Website: www.aulashack.com.br\n"
        "Email: contato@aulashack.com.br"
    )
    
    doc.add_heading('1.4 Propósito do Teste', 2)
    doc.add_paragraph(
        "Este teste de intrusão tem como objetivo avaliar a postura de segurança da "
        "aplicação web do Banco Digital Grana Fácil, identificando vulnerabilidades que "
        "possam ser exploradas por atacantes para comprometer a confidencialidade, "
        "integridade e disponibilidade dos sistemas e dados.\n\n"
        "O teste visa simular um ataque real por um atacante externo sem conhecimento "
        "prévio da aplicação (Black Box), permitindo ao banco identificar e corrigir "
        "falhas de segurança antes que sejam exploradas maliciosamente."
    )
    
    doc.add_page_break()

def add_section2(doc):
    """Seção 2: Contatos"""
    doc.add_heading('2. CONTATOS', 1)
    
    doc.add_heading('2.1 Equipe de Pentest (AulasHack)', 2)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Nome'
    header_cells[1].text = 'Função'
    header_cells[2].text = 'Email'
    header_cells[3].text = 'Telefone'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    pentest_team = [
        ('Carlos Mendes', 'Lead Pentester', 'carlos.mendes@aulashack.com.br', '(11) 98765-4321'),
        ('Ana Paula Silva', 'Senior Pentester', 'ana.silva@aulashack.com.br', '(11) 98765-4322'),
        ('Roberto Santos', 'Pentester', 'roberto.santos@aulashack.com.br', '(11) 98765-4323'),
        ('Mariana Costa', 'Project Manager', 'mariana.costa@aulashack.com.br', '(11) 98765-4320')
    ]
    
    for i, (nome, funcao, email, tel) in enumerate(pentest_team, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = nome
        row_cells[1].text = funcao
        row_cells[2].text = email
        row_cells[3].text = tel
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Contatos do Cliente (Banco Grana Fácil)', 2)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Nome'
    header_cells[1].text = 'Função'
    header_cells[2].text = 'Email'
    header_cells[3].text = 'Telefone'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    client_team = [
        ('João Silva', 'CISO', 'joao.silva@granafacil.com.br', '(11) 3000-1001'),
        ('Maria Oliveira', 'Gerente de Segurança', 'maria.oliveira@granafacil.com.br', '(11) 3000-1002'),
        ('Pedro Ferreira', 'Coordenador de TI', 'pedro.ferreira@granafacil.com.br', '(11) 3000-1003'),
        ('Luciana Souza', 'Desenvolvedora Líder', 'luciana.souza@granafacil.com.br', '(11) 3000-1004')
    ]
    
    for i, (nome, funcao, email, tel) in enumerate(client_team, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = nome
        row_cells[1].text = funcao
        row_cells[2].text = email
        row_cells[3].text = tel
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("IMPORTANTE: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run(
        "Todos os contatos devem estar disponíveis durante o período de teste para "
        "questões urgentes, especialmente em caso de incidentes ou descoberta de "
        "vulnerabilidades críticas."
    )
    
    doc.add_page_break()

def add_section3(doc):
    """Seção 3: Escopo do Teste"""
    doc.add_heading('3. ESCOPO DO TESTE', 1)
    
    doc.add_heading('3.1 Objetivos', 2)
    doc.add_paragraph("Os objetivos deste teste de intrusão incluem:")
    
    objectives = [
        'Identificar vulnerabilidades de segurança na aplicação web',
        'Avaliar a eficácia dos controles de autenticação e autorização',
        'Testar a resistência da aplicação contra ataques de injeção (SQL, Command)',
        'Verificar a implementação de melhores práticas de segurança web (OWASP Top 10)',
        'Avaliar a proteção de dados sensíveis e credenciais',
        'Identificar possíveis vetores de ataque para acesso não autorizado',
        'Testar a resposta da aplicação a tentativas de exploração',
        'Documentar todas as descobertas com evidências e recomendações de remediação'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('3.2 Alvos no Escopo', 2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("IMPORTANTE: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run(
        "Apenas os alvos listados abaixo estão autorizados para teste. Qualquer "
        "sistema fora desta lista NÃO deve ser testado."
    )
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Componente'
    header_cells[1].text = 'Detalhes'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    scope_data = [
        ('Ambiente', 'Ambiente de Homologação/Staging'),
        ('URL Principal', 'http://192.168.1.100:5000'),
        ('URL Alternativa', 'http://staging.granafacil.internal:5000'),
        ('Tipo de Aplicação', 'Aplicação Web - Banking Portal'),
        ('Tecnologia', 'Python/Flask, SQLite'),
        ('Funcionalidades', 'Login, Reset de Senha, Suporte'),
        ('Período de Teste', f'{date.today().strftime("%d/%m/%Y")} a {(date.today() + timedelta(days=14)).strftime("%d/%m/%Y")}')
    ]
    
    for i, (comp, det) in enumerate(scope_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = comp
        row_cells[1].text = det
    
    doc.add_paragraph()
    doc.add_paragraph("Áreas específicas a serem testadas:")
    
    areas = [
        'Mecanismo de autenticação (login)',
        'Funcionalidade de recuperação de senha',
        'Sistema de suporte técnico',
        'Validação e sanitização de entradas',
        'Controles de acesso e autorização',
        'Armazenamento de dados sensíveis',
        'Configurações de segurança do servidor',
        'Headers de segurança HTTP',
        'Páginas e diretórios ocultos'
    ]
    
    for area in areas:
        doc.add_paragraph(area, style='List Bullet')
    
    doc.add_heading('3.3 Fora do Escopo', 2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("CRÍTICO: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run(
        "Os seguintes alvos e atividades estão EXPRESSAMENTE PROIBIDOS e não devem "
        "ser testados sob nenhuma circunstância:"
    )
    
    doc.add_paragraph()
    
    out_of_scope = [
        'Ambiente de produção (www.granafacil.com.br)',
        'Banco de dados de produção',
        'Sistemas de terceiros integrados',
        'Infraestrutura de rede interna (switches, roteadores, firewalls)',
        'Servidores de email corporativo',
        'Sistemas de backup',
        'Qualquer sistema não explicitamente listado na seção 3.2',
        'Testes físicos nas instalações do banco',
        'Engenharia social contra funcionários',
        'Ataques de negação de serviço (DoS/DDoS)'
    ]
    
    for item in out_of_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()

def add_section4(doc):
    """Seção 4: Tipo e Metodologia"""
    doc.add_heading('4. TIPO E METODOLOGIA', 1)
    
    doc.add_heading('4.1 Tipo de Teste: Black Box', 2)
    
    doc.add_paragraph(
        "Este é um teste Black Box (caixa preta), onde a equipe de pentest não possui "
        "conhecimento prévio sobre a arquitetura interna, código-fonte ou configurações "
        "da aplicação. Esta abordagem simula um atacante externo real que tenta "
        "comprometer o sistema sem informações privilegiadas."
    )
    
    doc.add_paragraph()
    doc.add_paragraph("Características do teste Black Box:")
    
    characteristics = [
        'Sem acesso ao código-fonte',
        'Sem credenciais iniciais',
        'Sem conhecimento da arquitetura interna',
        'Sem informações sobre tecnologias utilizadas (será descoberto)',
        'Perspectiva de um atacante externo',
        'Foco em vulnerabilidades exploráveis externamente'
    ]
    
    for char in characteristics:
        doc.add_paragraph(char, style='List Bullet')
    
    doc.add_heading('4.2 Metodologia', 2)
    
    doc.add_paragraph(
        "O teste seguirá metodologias reconhecidas internacionalmente:"
    )
    
    doc.add_paragraph()
    doc.add_paragraph("PTES (Penetration Testing Execution Standard):")
    
    ptes_phases = [
        'Planejamento e Preparação',
        'Reconhecimento e Descoberta (Information Gathering)',
        'Análise de Vulnerabilidades (Vulnerability Assessment)',
        'Exploração (Exploitation)',
        'Pós-Exploração (Post-Exploitation)',
        'Documentação e Relatório'
    ]
    
    for phase in ptes_phases:
        doc.add_paragraph(phase, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("Frameworks de referência:")
    
    frameworks = [
        'OWASP Testing Guide v4.2',
        'OWASP Top 10 2021',
        'NIST SP 800-115 (Technical Guide to Information Security Testing)',
        'CWE Top 25 Most Dangerous Software Weaknesses',
        'PTES Technical Guidelines'
    ]
    
    for fw in frameworks:
        doc.add_paragraph(fw, style='List Bullet')
    
    doc.add_heading('4.3 Ferramentas Autorizadas', 2)
    
    doc.add_paragraph(
        "As seguintes categorias de ferramentas poderão ser utilizadas:"
    )
    
    tools = [
        'Scanners de vulnerabilidades web (Burp Suite, OWASP ZAP, Nikto)',
        'Ferramentas de análise de rede (Nmap, Wireshark)',
        'Ferramentas de fuzzing (Gobuster, Dirsearch, ffuf)',
        'Ferramentas de exploração (SQLMap, Metasploit)',
        'Scripts customizados em Python',
        'Ferramentas de análise de tráfego',
        'Decodificadores e analisadores (CyberChef)',
        'Ferramentas de força bruta controlada (Hydra, com limitações)'
    ]
    
    for tool in tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_page_break()

def add_section5(doc):
    """Seção 5: Cronograma"""
    doc.add_heading('5. CRONOGRAMA', 1)
    
    hoje = date.today()
    
    doc.add_paragraph(
        "O teste de intrusão será realizado durante um período de 14 dias úteis, "
        "distribuído nas seguintes fases:"
    )
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Fase'
    header_cells[1].text = 'Atividades'
    header_cells[2].text = 'Duração'
    header_cells[3].text = 'Datas'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    schedule_data = [
        ('1. Kick-off', 'Reunião inicial, alinhamentos', '1 dia', f'{hoje.strftime("%d/%m")}'),
        ('2. Reconhecimento', 'Descoberta, mapeamento', '2 dias', f'{(hoje + timedelta(days=1)).strftime("%d/%m")} - {(hoje + timedelta(days=2)).strftime("%d/%m")}'),
        ('3. Análise', 'Scanning, identificação', '2 dias', f'{(hoje + timedelta(days=3)).strftime("%d/%m")} - {(hoje + timedelta(days=4)).strftime("%d/%m")}'),
        ('4. Exploração', 'Testes de exploração', '4 dias', f'{(hoje + timedelta(days=5)).strftime("%d/%m")} - {(hoje + timedelta(days=8)).strftime("%d/%m")}'),
        ('5. Validação', 'Confirmação, reteste', '2 dias', f'{(hoje + timedelta(days=9)).strftime("%d/%m")} - {(hoje + timedelta(days=10)).strftime("%d/%m")}'),
        ('6. Documentação', 'Elaboração do relatório', '2 dias', f'{(hoje + timedelta(days=11)).strftime("%d/%m")} - {(hoje + timedelta(days=12)).strftime("%d/%m")}'),
        ('7. Apresentação', 'Reunião de encerramento', '1 dia', f'{(hoje + timedelta(days=13)).strftime("%d/%m")}')
    ]
    
    for i, (fase, ativ, dur, datas) in enumerate(schedule_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = fase
        row_cells[1].text = ativ
        row_cells[2].text = dur
        row_cells[3].text = datas
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Nota: ")
    run.bold = True
    p.add_run(
        "O cronograma pode ser ajustado conforme necessário mediante acordo entre "
        "as partes. Atrasos causados por indisponibilidade do ambiente ou descoberta "
        "de vulnerabilidades críticas que requerem remediação imediata podem impactar "
        "o cronograma."
    )
    
    doc.add_page_break()

def add_section6(doc):
    """Seção 6: Limitações e Restrições"""
    doc.add_heading('6. LIMITAÇÕES E RESTRIÇÕES', 1)
    
    doc.add_heading('6.1 Restrições de Horário', 2)
    
    p = doc.add_paragraph()
    run = p.add_run("IMPORTANTE - Horários de Teste Permitidos:\n\n")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph(
        "Para minimizar impactos no ambiente e evitar interferência com atividades "
        "de desenvolvimento e testes regulares, os seguintes horários foram estabelecidos:"
    )
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Período'
    header_cells[1].text = 'Status'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    time_data = [
        ('Segunda a Sexta: 10:00 - 17:00', '✓ PERMITIDO (horário preferencial)'),
        ('Segunda a Sexta: 08:00 - 10:00', '✗ PROIBIDO (horário de pico)'),
        ('Segunda a Sexta: 17:00 - 20:00', '✗ PROIBIDO (horário de pico)')
    ]
    
    for i, (periodo, status) in enumerate(time_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = periodo
        row_cells[1].text = status
        if '✗' in status:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
        else:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 150, 0)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Justificativa: ")
    run.bold = True
    p.add_run(
        "Os horários de 08:00-10:00 e 17:00-20:00 são considerados horários de pico, "
        "onde a equipe de desenvolvimento está mais ativa e testes automatizados "
        "regulares são executados. Testes de pentest durante estes períodos podem "
        "causar confusão, falsos positivos em monitoramento e possíveis conflitos "
        "com atividades legítimas."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Exceções: ")
    run.bold = True
    p.add_run(
        "Testes fora dos horários permitidos podem ser autorizados mediante aprovação "
        "prévia por escrito do CISO (João Silva) com pelo menos 24 horas de antecedência."
    )
    
    doc.add_heading('6.2 Atividades Proibidas', 2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("As seguintes atividades são ESTRITAMENTE PROIBIDAS:\n\n")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    prohibited = [
        'Ataques de Negação de Serviço (DoS/DDoS)',
        'Flooding ou sobrecarga intencional de recursos',
        'Modificação ou exclusão de dados em produção',
        'Testes em ambiente de produção',
        'Engenharia social contra funcionários',
        'Ataques físicos às instalações',
        'Phishing ou spear-phishing',
        'Exploração de vulnerabilidades zero-day sem aprovação prévia',
        'Instalação de backdoors permanentes',
        'Exfiltração de dados reais de clientes',
        'Testes automatizados agressivos durante horários de pico',
        'Compartilhamento de descobertas com terceiros',
        'Uso de ferramentas não aprovadas pelo cliente'
    ]
    
    for item in prohibited:
        doc.add_paragraph(f"✗ {item}", style='List Bullet')
    
    doc.add_heading('6.3 Limitações Técnicas', 2)
    
    doc.add_paragraph(
        "As seguintes limitações técnicas devem ser observadas:"
    )
    
    limitations = [
        'Taxa máxima de requisições: 100 req/segundo (para evitar DoS acidental)',
        'Tamanho máximo de payload: 10MB',
        'Threads simultâneos: Máximo de 10 threads em ferramentas automatizadas',
        'Tentativas de força bruta: Máximo de 1000 tentativas por conta',
        'Tempo máximo de exploit: 30 segundos por tentativa',
        'Largura de banda: Uso moderado, sem saturar a rede',
        'Armazenamento: Dados coletados devem ser armazenados criptografados'
    ]
    
    for limit in limitations:
        doc.add_paragraph(limit, style='List Bullet')
    
    doc.add_page_break()

def add_section7(doc):
    """Seção 7: Regras de Engajamento"""
    doc.add_heading('7. REGRAS DE ENGAJAMENTO', 1)
    
    doc.add_heading('7.1 Comunicação', 2)
    
    doc.add_paragraph("Canais de comunicação autorizados:")
    
    comm_channels = [
        'Email: contato@aulashack.com.br (comunicações formais)',
        'Email urgente: urgent@aulashack.com.br (vulnerabilidades críticas)',
        'Telefone: (11) 98765-4320 (emergências 24/7)',
        'Slack: Canal #pentest-granafacil (comunicações rápidas)',
        'Reuniões: Calls diárias às 09:00 e 18:00 (15 minutos)'
    ]
    
    for channel in comm_channels:
        doc.add_paragraph(channel, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Protocolo de Comunicação:\n")
    run.bold = True
    
    protocols = [
        'Updates diários ao final do dia com resumo das atividades',
        'Notificação IMEDIATA de vulnerabilidades críticas (CVSS ≥ 9.0)',
        'Notificação em até 4 horas de vulnerabilidades altas (CVSS 7.0-8.9)',
        'Notificação em até 24 horas de vulnerabilidades médias',
        'Todas as comunicações devem ser criptografadas (PGP ou S/MIME)',
        'Nenhuma informação sensível por SMS ou chamadas não criptografadas'
    ]
    
    for protocol in protocols:
        doc.add_paragraph(protocol, style='List Bullet')
    
    doc.add_heading('7.2 Procedimentos de Emergência', 2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Em caso de EMERGÊNCIA (crash do sistema, dados corrompidos, etc.):\n\n")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    emergency_steps = [
        'PARE IMEDIATAMENTE todos os testes',
        'Documente o que estava sendo testado no momento',
        'Notifique IMEDIATAMENTE o CISO (João Silva) via telefone: (11) 3000-1001',
        'Envie email para: urgent@granafacil.com.br',
        'Aguarde instruções antes de retomar testes',
        'Preencha relatório de incidente em até 1 hora'
    ]
    
    for i, step in enumerate(emergency_steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Contatos de Emergência 24/7:\n")
    run.bold = True
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Nome'
    header_cells[1].text = 'Função'
    header_cells[2].text = 'Telefone'
    
    emergency_contacts = [
        ('João Silva', 'CISO', '(11) 3000-1001'),
        ('Maria Oliveira', 'Gerente de Segurança', '(11) 3000-1002')
    ]
    
    for i, (nome, funcao, tel) in enumerate(emergency_contacts, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = nome
        row_cells[1].text = funcao
        row_cells[2].text = tel
    
    doc.add_heading('7.3 Confidencialidade', 2)
    
    doc.add_paragraph(
        "Todas as informações obtidas durante o teste são estritamente confidenciais:"
    )
    
    confidentiality = [
        'NDA (Non-Disclosure Agreement) assinado por todos os membros da equipe',
        'Dados coletados armazenados em dispositivos criptografados',
        'Acesso restrito apenas à equipe de pentest autorizada',
        'Destruição segura de todos os dados ao final do projeto',
        'Proibição de discussão pública das descobertas',
        'Relatórios marcados como CONFIDENCIAL',
        'Compartilhamento apenas com partes autorizadas',
        'Retenção de dados por no máximo 90 dias após entrega do relatório'
    ]
    
    for item in confidentiality:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.4 Regras de Evidência', 2)
    
    doc.add_paragraph(
        "Coleta e tratamento de evidências:"
    )
    
    evidence_rules = [
        'Screenshots devem incluir data/hora e URL completa',
        'Comandos executados devem ser documentados com output completo',
        'Payloads de exploração devem ser salvos para reprodução',
        'Logs de ferramentas devem ser preservados',
        'Evidências devem ser numeradas e referenciadas no relatório',
        'Dados sensíveis em screenshots devem ser censurados',
        'Backup de todas as evidências em local seguro',
        'Chain of custody mantida para todas as evidências'
    ]
    
    for rule in evidence_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()

def add_section8(doc):
    """Seção 8: Entregáveis"""
    doc.add_heading('8. ENTREGÁVEIS', 1)
    
    doc.add_paragraph(
        "Ao final do teste de intrusão, os seguintes entregáveis serão fornecidos:"
    )
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Entregável'
    header_cells[1].text = 'Formato'
    header_cells[2].text = 'Prazo'
    
    deliverables = [
        ('Relatório Executivo', 'PDF', 'D+3 dias'),
        ('Relatório Técnico Completo', 'DOCX + PDF', 'D+5 dias'),
        ('Evidências Técnicas', 'ZIP criptografado', 'D+5 dias'),
        ('Apresentação', 'PPTX + PDF', 'D+5 dias'),
        ('Reunião de Apresentação', 'Presencial/Online', 'D+7 dias')
    ]
    
    for i, (entregavel, formato, prazo) in enumerate(deliverables, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = entregavel
        row_cells[1].text = formato
        row_cells[2].text = prazo
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Nota: ")
    run.bold = True
    p.add_run("D = Data de conclusão dos testes (último dia da fase de validação)")
    
    doc.add_paragraph()
    
    doc.add_heading('8.1 Conteúdo do Relatório', 2)
    
    report_content = [
        'Sumário Executivo (para gestão)',
        'Sumário Técnico',
        'Metodologia utilizada',
        'Escopo e limitações',
        'Descobertas detalhadas (uma por vulnerabilidade)',
        'Classificação de severidade (CVSS 3.1)',
        'Evidências (screenshots, comandos, outputs)',
        'Recomendações de remediação priorizadas',
        'Referências técnicas (CWE, OWASP, CVE)',
        'Plano de reteste sugerido'
    ]
    
    for item in report_content:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()

def add_section9(doc):
    """Seção 9: Critérios de Aceitação"""
    doc.add_heading('9. CRITÉRIOS DE ACEITAÇÃO', 1)
    
    doc.add_paragraph(
        "O teste de intrusão será considerado bem-sucedido se os seguintes "
        "critérios forem atendidos:"
    )
    
    doc.add_paragraph()
    
    acceptance_criteria = [
        'Todos os componentes no escopo foram testados adequadamente',
        'Metodologia PTES foi seguida em todas as fases',
        'Todas as vulnerabilidades identificadas foram documentadas',
        'Evidências suficientes foram coletadas para cada descoberta',
        'Classificação CVSS aplicada a todas as vulnerabilidades',
        'Recomendações de remediação fornecidas para cada descoberta',
        'Relatório entregue dentro do prazo estabelecido',
        'Apresentação realizada com stakeholders',
        'Nenhum dano foi causado ao ambiente',
        'Todas as regras de engajamento foram respeitadas',
        'Comunicação adequada mantida durante todo o processo',
        'Dados coletados foram tratados com confidencialidade'
    ]
    
    for criteria in acceptance_criteria:
        doc.add_paragraph(f"✓ {criteria}", style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('9.1 Indicadores de Qualidade', 2)
    
    quality_indicators = [
        'Taxa de falsos positivos < 5%',
        'Reprodutibilidade das vulnerabilidades: 100%',
        'Clareza das evidências: Alta',
        'Aplicabilidade das recomendações: Alta',
        'Satisfação do cliente: ≥ 4/5',
        'Aderência ao cronograma: ± 2 dias'
    ]
    
    for indicator in quality_indicators:
        doc.add_paragraph(indicator, style='List Bullet')
    
    doc.add_page_break()

def add_section10(doc):
    """Seção 10: Aprovações"""
    doc.add_heading('10. APROVAÇÕES', 1)
    
    doc.add_paragraph(
        "Este Plano de Teste foi revisado e aprovado pelas seguintes partes interessadas:"
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Cliente
    p = doc.add_paragraph()
    run = p.add_run("BANCO DIGITAL GRANA FÁCIL S.A.\n\n")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph("_" * 50)
    doc.add_paragraph("João Silva")
    doc.add_paragraph("Chief Information Security Officer (CISO)")
    doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Empresa de pentest
    p = doc.add_paragraph()
    run = p.add_run("AULASHACK SECURITY CONSULTING\n\n")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph("_" * 50)
    doc.add_paragraph("Carlos Mendes")
    doc.add_paragraph("Lead Pentester")
    doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph("_" * 50)
    doc.add_paragraph("Mariana Costa")
    doc.add_paragraph("Project Manager")
    doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Nota: ")
    run.bold = True
    p.add_run(
        "Este documento deve ser assinado por todas as partes antes do início "
        "dos testes. Qualquer alteração neste plano requer aprovação mútua por escrito."
    )

def main():
    """Função principal"""
    print("[*] Gerando Plano de Teste...")
    
    doc = Document()
    
    # Capa
    print("    [+] Adicionando capa...")
    add_cover_page(doc)
    
    # Informações do documento
    print("    [+] Adicionando informações do documento...")
    add_document_info(doc)
    
    # Sumário
    print("    [+] Adicionando sumário...")
    add_toc(doc)
    
    # Seções
    print("    [+] Adicionando seção 1 - Informações do Projeto...")
    add_section1(doc)
    
    print("    [+] Adicionando seção 2 - Contatos...")
    add_section2(doc)
    
    print("    [+] Adicionando seção 3 - Escopo...")
    add_section3(doc)
    
    print("    [+] Adicionando seção 4 - Tipo e Metodologia...")
    add_section4(doc)
    
    print("    [+] Adicionando seção 5 - Cronograma...")
    add_section5(doc)
    
    print("    [+] Adicionando seção 6 - Limitações e Restrições...")
    add_section6(doc)
    
    print("    [+] Adicionando seção 7 - Regras de Engajamento...")
    add_section7(doc)
    
    print("    [+] Adicionando seção 8 - Entregáveis...")
    add_section8(doc)
    
    print("    [+] Adicionando seção 9 - Critérios de Aceitação...")
    add_section9(doc)
    
    print("    [+] Adicionando seção 10 - Aprovações...")
    add_section10(doc)
    
    # Salvar
    output_path = '/home/claude/Plano_de_Teste_Banco_Grana_Facil.docx'
    doc.save(output_path)
    
    print(f"\n[✓] Plano de Teste gerado com sucesso!")
    print(f"[✓] Arquivo: {output_path}\n")

if __name__ == '__main__':
    main()
